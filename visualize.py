import gradio as gr

import os
import pandas as pd
import docx

import re
from ftfy import fix_text

from HSmodule import *
from source.FaissSearch import *

from source.BloomDetection import *
from source.Preprocessor import *
from source.SimHashDetection import *
from source.minHashDetection import *
from sentence_transformers import SentenceTransformer


def read_file(filepath):
    #tach file extension
    ext = os.path.splitext(filepath)[1].lower()
    # tach du lieu tu file
    if ext == ".docx":
      doc = docx.Document(filepath)
      paragraphs = [(p.text.strip()) for p in doc.paragraphs if p.text.strip()]
    elif ext == ".txt":
      with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
      paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    elif ext == ".csv":
      table = pd.read_csv(filepath)
      data_cols = ["content", "text", "paragraph"]
      for col in data_cols:
        if col in table.columns:
           paragraphs = table[col].dropna().astype(str).tolist()
           break
    else:
        raise ValueError("Kh√¥ng t√¨m th·∫•y c·ªôt 'text' / 'content' / 'paragraph' trong file CSV ho·∫∑c file kh√¥ng h·ª£p l·ªá.")
    return paragraphs

modelSimHash = SimHashDetection()
modelMinHash = MinHashDetection()
modelBloomFaiss = BloomDetection()

def run_SimHash(paragraphs : list[str]) -> list[list[VectorRecord]]:
    doing = modelSimHash
    return doing.detect(paragraphs)


def run_Bloom_Sim_Faiss(paragraphs : list[str]) -> list[list[VectorRecord]]:
    doing = modelBloomFaiss
    return doing.detect(paragraphs)


def run_Min(paragraphs : list[str]) -> list[list[VectorRecord]]:
    doing = modelMinHash
    return doing.detect(paragraphs)


def representative_texts(l : list[list[VectorRecord]], p : list[str]) -> dict():
    text = {min(x.id for x in group) : max([p[x.id] for x in group], key=len) for group in l}
    return dict(sorted(text.items()))


def duplication_text(filepath : str, method : str):

    # tach du lieu
    paragraphs = read_file(filepath)
    if method == "SimHash (Semantic)":
      ans = run_SimHash(paragraphs)
    elif method == "Bloom + Faiss (Semantic)":
      ans = run_Bloom_Sim_Faiss(paragraphs)
    elif method == "MinHash (Syntax)":
      ans = run_Min(paragraphs)

    # loc van ban dai dien
    filted_text = representative_texts(ans, paragraphs)




    #hightlight
    # m√†u ƒë·ªÉ highlight - 17 m√†u
    from docx.enum.text import WD_COLOR_INDEX
    colors = [
        None,  # 0 ‚Üí kh√¥ng highlight
        WD_COLOR_INDEX.BLACK,
        WD_COLOR_INDEX.BLUE,
        WD_COLOR_INDEX.BRIGHT_GREEN,
        WD_COLOR_INDEX.DARK_BLUE,
        WD_COLOR_INDEX.DARK_RED,
        WD_COLOR_INDEX.DARK_YELLOW,
        WD_COLOR_INDEX.GRAY_25,
        WD_COLOR_INDEX.GRAY_50,
        WD_COLOR_INDEX.GREEN,
        WD_COLOR_INDEX.PINK,
        WD_COLOR_INDEX.RED,
        WD_COLOR_INDEX.TEAL,
        WD_COLOR_INDEX.TURQUOISE,
        WD_COLOR_INDEX.VIOLET,
        WD_COLOR_INDEX.WHITE,
        WD_COLOR_INDEX.YELLOW,
    ]

    # t·∫°o dictionraty ƒë·ªÉ truy xu·∫•t x·ª≠ l√Ω d·ªØ li·ªáu
    group_id = {x.id: (0 if len(group) == 1 else group_index) for group_index, group in enumerate(ans) for x in group}

    doc = docx.Document()
    doc.add_heading("Duplicate Text Highlighting Result")
    # Th√™m t·ª´ng paragraph t·ª´ filted_text
    for para in filted_text.values():
        doc.add_paragraph(para)

    # L∆∞u file
    result = "result.docx"
    doc.save(result)



    # in ket qua
    # thanh cuon
    css = """
    <style>
    .simple-scroll-content {
        max-height: 500px;
        overflow-y: auto;
        padding: 20px;
        border: 1px solid #444;
        border-radius: 10px;
        background: #000;
        color: white;
        margin-top: 10px;
    }
    </style>
    """


    html1 = css + """
    <div class="simple-scroll-content">
    """

    for id, group in enumerate(ans):
      html1 += f"<h3 style='color:#0af'>Group {id} ({len(group)} items)</h3><ul>"
      for para in group:
        html1 += f"<li><b>Paragraph {para.id}</b>: {paragraphs[para.id][:300]}...</li>"
      html1 += "</ul>"

    html1 += "</div>"


    # van ban sau khi loc
    html2 = css + """
    <div class="simple-scroll-content">
    """

    for para in filted_text.values():
      html2 += f"<h3 style='color:#0af'>{para}</h3>"

    html2 += "</div>"

    return html1, html2, result

# d√πng gradio ƒë·ªÉ t·∫°o giao di·ªán demo
with gr.Blocks(title="Duplicate Text Detector") as demo:
    gr.Markdown("## üß© Duplicate Text Detector")

    with gr.Row():
        with gr.Column():
            file_input = gr.File(label="Upload document file")
            method_choice = gr.Radio(
                ["SimHash (Semantic)", "Bloom + Faiss (Semantic)", "MinHash (Syntax)"],
                label="Choose method :",
                value="SimHash (Semantic)"
            )
            submit_btn = gr.Button("Submit", variant="primary")

    with gr.Row():
        with gr.Column():
            with gr.Accordion("üìä Duplication Text Group Result", open=False) as accordion_1:
              html_output_1 = gr.HTML()

        with gr.Column():
            with gr.Accordion("üìù Text after removing duplication", open=False) as accordion_2:
              html_output_2 = gr.HTML()

    with gr.Row():
        file_output = gr.File(label="Download result")

    submit_btn.click(
        fn=duplication_text,
        inputs=[file_input, method_choice],
        outputs=[html_output_1, html_output_2, file_output]
    )
    
port = int(os.environ.get("PORT", 10000))  # Render cung c·∫•p PORT, fallback 10000 khi ch·∫°y local
demo.launch(server_name="0.0.0.0", server_port=port)
